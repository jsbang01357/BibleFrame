#!/usr/bin/env python3
"""BibleFrame 73권 본문을 TXT·JSON·JSONL·DOCX·PDF 다운로드로 만든다."""

from __future__ import annotations

import json
import zipfile
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer
from reportlab.platypus.tableofcontents import TableOfContents


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "site" / "bible.json"
DOWNLOADS = ROOT / "site" / "downloads"
STEM = "bibleframe-ko-catholic-73"
TITLE = "BibleFrame 한국어 가톨릭 성경 73권"
SUBTITLE = "Public Domain WEB-C 기반 · Qwen3-Next 80B 기계 번역·교정"
NOTICE = "가톨릭 정경 기반 비공인 기계 번역 초안"
DATA_LICENSE = "CC0-1.0"
SOURCE_URL = "https://ebible.org/bible/details.php?id=eng-web-c"
PROJECT_URL = "https://github.com/jsbang01357/BibleFrame"
BODY_FONT = "NanumGothic"
HEADING_FONT = "NanumGothic"
PDF_FONT_CANDIDATES = (
    Path.home() / "Library/Fonts/NanumGothic.ttf",
    Path("/Library/Fonts/NanumGothic.ttf"),
    Path("/usr/share/fonts/truetype/nanum/NanumGothic.ttf"),
    Path("/usr/share/fonts/opentype/nanum/NanumGothic.ttf"),
)
FIXED_ZIP_TIME = (1980, 1, 1, 0, 0, 0)


class BiblePdfTemplate(SimpleDocTemplate):
    """책 제목을 PDF 북마크와 자동 목차 항목으로 연결한다."""

    def afterFlowable(self, flowable) -> None:
        if not isinstance(flowable, Paragraph) or flowable.style.name != "BibleBook":
            return
        title = flowable.getPlainText()
        key = getattr(flowable, "_bookmark_name", None)
        if not key:
            return
        self.canv.bookmarkPage(key)
        self.canv.addOutlineEntry(title, key, level=0, closed=False)
        self.notify("TOCEntry", (0, title, self.page, key))


def load_payload() -> dict[str, object]:
    payload = json.loads(SOURCE.read_text(encoding="utf-8"))
    if payload["meta"]["books"] != 73 or payload["meta"]["verses"] != 35_379:
        raise ValueError("검증된 73권 사이트 데이터가 아님")
    return payload


def group_verses(verses: list[dict[str, object]]) -> dict[tuple[str, int], list[dict[str, object]]]:
    chapters: dict[tuple[str, int], list[dict[str, object]]] = defaultdict(list)
    for verse in verses:
        chapters[(str(verse["code"]), int(verse["chapter"]))].append(verse)
    return chapters


def write_text(payload: dict[str, object]) -> None:
    books = payload["books"]
    verses = payload["verses"]
    chapters = group_verses(verses)
    lines = [
        TITLE,
        SUBTITLE,
        NOTICE,
        f"데이터 라이선스: {DATA_LICENSE}",
        f"영어 원문: {SOURCE_URL}",
        f"프로젝트: {PROJECT_URL}",
        "",
    ]
    for book in books:
        code = str(book["code"])
        lines.extend((f"# {book['name']} ({book['english']})", ""))
        chapter_numbers = sorted(chapter for book_code, chapter in chapters if book_code == code)
        for chapter in chapter_numbers:
            lines.append(f"## {book['name']} {chapter}장")
            lines.extend(f"{item['verse']} {item['text']}" for item in chapters[(code, chapter)])
            lines.append("")
    (DOWNLOADS / f"{STEM}.txt").write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def write_json(payload: dict[str, object]) -> None:
    (DOWNLOADS / f"{STEM}.json").write_text(
        json.dumps(payload, ensure_ascii=False, separators=(",", ":")) + "\n", encoding="utf-8"
    )


def write_jsonl(payload: dict[str, object]) -> None:
    path = DOWNLOADS / f"{STEM}.jsonl"
    with path.open("w", encoding="utf-8") as stream:
        for item in payload["verses"]:
            record = {
                "id": item["id"],
                "code": item["code"],
                "book": item["book"],
                "short": item["short"],
                "english": item["english"],
                "testament": item["testament"],
                "chapter": item["chapter"],
                "verse": item["verse"],
                "text": item["text"],
                "language": "ko",
                "translation": TITLE,
                "translation_status": "machine_reviewed_draft",
                "translation_model": "Qwen3-Next-80B-A3B-Instruct",
                "source": "World English Bible (Catholic)",
                "source_license": "Public Domain",
                "data_license": DATA_LICENSE,
            }
            stream.write(json.dumps(record, ensure_ascii=False, separators=(",", ":")) + "\n")


def set_run_font(run, name: str, size: float, *, bold: bool | None = None,
                 color: RGBColor | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def configure_style(style, *, font: str, size: float, color: RGBColor,
                    before: float, after: float, line_spacing: float,
                    bold: bool = False) -> None:
    style.font.name = font
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line_spacing


def normalize_office_zip(path: Path) -> None:
    """DOCX ZIP 메타데이터를 고정해 같은 입력에서 같은 파일을 만든다."""
    with zipfile.ZipFile(path, "r") as source:
        entries = [(name, source.read(name)) for name in source.namelist()]
    temporary = path.with_suffix(f"{path.suffix}.tmp")
    with zipfile.ZipFile(temporary, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as target:
        for name, data in entries:
            info = zipfile.ZipInfo(name, FIXED_ZIP_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.create_system = 3
            mode = 0o40755 if name.endswith("/") else 0o100644
            info.external_attr = mode << 16
            target.writestr(info, data)
    temporary.replace(path)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("BibleFrame  ·  ")
    set_run_font(run, HEADING_FONT, 8.5, color=RGBColor(110, 104, 94))
    page_run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_run._r.extend((fld_begin, instr, fld_end))
    set_run_font(page_run, HEADING_FONT, 8.5, color=RGBColor(110, 104, 94))


def write_docx(payload: dict[str, object]) -> None:
    books = payload["books"]
    verses = payload["verses"]
    chapters = group_verses(verses)
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = document.styles["Normal"]
    configure_style(normal, font=BODY_FONT, size=11, color=RGBColor(25, 23, 19),
                    before=0, after=6, line_spacing=1.25)
    h1 = document.styles["Heading 1"]
    configure_style(h1, font=HEADING_FONT, size=16, color=RGBColor(46, 116, 181),
                    before=18, after=10, line_spacing=1.0, bold=True)
    h1.paragraph_format.page_break_before = True
    h1.paragraph_format.keep_with_next = True
    h2 = document.styles["Heading 2"]
    configure_style(h2, font=HEADING_FONT, size=13, color=RGBColor(46, 116, 181),
                    before=14, after=7, line_spacing=1.0, bold=True)
    h2.paragraph_format.keep_with_next = True
    h3 = document.styles["Heading 3"]
    configure_style(h3, font=HEADING_FONT, size=12, color=RGBColor(31, 77, 120),
                    before=10, after=5, line_spacing=1.0, bold=True)

    core = document.core_properties
    core.title = TITLE
    core.subject = NOTICE
    core.author = "BibleFrame"
    core.keywords = "가톨릭 성경, 73권, 오픈 데이터, CC0, Qwen3-Next 80B"
    core.comments = f"{SUBTITLE} · {PROJECT_URL}"
    core.created = datetime(2026, 7, 13)
    core.modified = datetime(2026, 7, 13)

    header = section.header.paragraphs[0]
    header.text = "BibleFrame  |  한국어 가톨릭 성경 73권"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        set_run_font(run, HEADING_FONT, 8.5, color=RGBColor(110, 104, 94))
    add_page_number(section.footer.paragraphs[0])

    cover = document.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(160)
    cover.paragraph_format.space_after = Pt(12)
    set_run_font(cover.add_run(TITLE), HEADING_FONT, 28, bold=True, color=RGBColor(32, 55, 72))
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    set_run_font(subtitle.add_run(SUBTITLE), HEADING_FONT, 13.5, color=RGBColor(43, 81, 99))
    for text, bold in (
        (NOTICE, True),
        ("한국천주교주교회의 공용 번역본이 아니며 전례용·교리 판정용이 아닙니다.", False),
        (f"한국어 본문·데이터 {DATA_LICENSE} · 코드 MIT", False),
        (PROJECT_URL, False),
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(7)
        set_run_font(paragraph.add_run(text), BODY_FONT, 10.5, bold=bold,
                     color=RGBColor(80, 80, 80))

    for book in books:
        code = str(book["code"])
        book_heading = document.add_paragraph(style="Heading 1")
        book_heading.add_run(f"{book['name']}  |  {book['english']}")
        chapter_numbers = sorted(chapter for book_code, chapter in chapters if book_code == code)
        for chapter in chapter_numbers:
            chapter_heading = document.add_paragraph(f"{book['name']} {chapter}장", style="Heading 2")
            chapter_heading.paragraph_format.keep_with_next = True
            body = document.add_paragraph()
            body.paragraph_format.widow_control = True
            for index, item in enumerate(chapters[(code, chapter)]):
                if index:
                    body.add_run(" ")
                number = body.add_run(str(item["verse"]))
                set_run_font(number, HEADING_FONT, 7.5, bold=True, color=RGBColor(156, 109, 45))
                number.font.superscript = True
                text = body.add_run(f" {item['text']}")
                set_run_font(text, BODY_FONT, 11, color=RGBColor(25, 23, 19))

    path = DOWNLOADS / f"{STEM}.docx"
    document.save(path)
    normalize_office_zip(path)


def write_pdf(payload: dict[str, object]) -> None:
    books = payload["books"]
    verses = payload["verses"]
    chapters = group_verses(verses)
    font_path = next((path for path in PDF_FONT_CANDIDATES if path.exists()), None)
    if font_path is None:
        raise FileNotFoundError("PDF 생성에 필요한 NanumGothic.ttf를 찾지 못함")
    pdfmetrics.registerFont(TTFont("BibleKorean", str(font_path)))

    path = DOWNLOADS / f"{STEM}.pdf"
    document = BiblePdfTemplate(
        str(path), pagesize=letter,
        rightMargin=0.72 * inch, leftMargin=0.72 * inch,
        topMargin=0.76 * inch, bottomMargin=0.72 * inch,
        title=TITLE, author="BibleFrame", subject=NOTICE,
        creator="BibleFrame build_downloads.py",
        pageCompression=1, invariant=1,
    )
    title_style = ParagraphStyle(
        "CoverTitle", fontName="BibleKorean", fontSize=25, leading=34,
        textColor=colors.HexColor("#203748"), alignment=TA_CENTER, spaceAfter=12,
    )
    subtitle_style = ParagraphStyle(
        "CoverSubtitle", fontName="BibleKorean", fontSize=12.5, leading=19,
        textColor=colors.HexColor("#2B5163"), alignment=TA_CENTER, spaceAfter=28,
    )
    cover_style = ParagraphStyle(
        "CoverInfo", fontName="BibleKorean", fontSize=9.5, leading=16,
        textColor=colors.HexColor("#555555"), alignment=TA_CENTER, spaceAfter=7,
    )
    book_style = ParagraphStyle(
        "BibleBook", fontName="BibleKorean", fontSize=17, leading=24,
        textColor=colors.HexColor("#2E74B5"), spaceAfter=12, keepWithNext=True,
    )
    chapter_style = ParagraphStyle(
        "Chapter", fontName="BibleKorean", fontSize=12.5, leading=18,
        textColor=colors.HexColor("#2E74B5"), spaceBefore=10, spaceAfter=6,
        keepWithNext=True,
    )
    body_style = ParagraphStyle(
        "BibleBody", fontName="BibleKorean", fontSize=9.8, leading=16,
        textColor=colors.HexColor("#191713"), alignment=TA_LEFT,
        spaceAfter=5, wordWrap="CJK", splitLongWords=True,
    )
    toc_title_style = ParagraphStyle(
        "TocTitle", fontName="BibleKorean", fontSize=21, leading=28,
        textColor=colors.HexColor("#203748"), alignment=TA_CENTER,
        spaceBefore=8, spaceAfter=22,
    )
    toc = TableOfContents()
    toc.levelStyles = [ParagraphStyle(
        "TocBook", fontName="BibleKorean", fontSize=10, leading=15,
        textColor=colors.HexColor("#2B5163"), leftIndent=18,
        firstLineIndent=-12, spaceBefore=1, spaceAfter=1,
    )]

    story = [
        Spacer(1, 2.0 * inch),
        Paragraph(escape(TITLE), title_style),
        Paragraph(escape(SUBTITLE), subtitle_style),
        Paragraph(f"<b>{escape(NOTICE)}</b>", cover_style),
        Paragraph("한국천주교주교회의 공용 번역본이 아니며 전례용·교리 판정용이 아닙니다.", cover_style),
        Paragraph(f"한국어 본문·데이터 {DATA_LICENSE} · 코드 MIT", cover_style),
        Paragraph(escape(PROJECT_URL), cover_style),
        PageBreak(),
        Paragraph("목차", toc_title_style),
        toc,
        PageBreak(),
    ]
    first_book = True
    for book in books:
        if not first_book:
            story.append(PageBreak())
        first_book = False
        code = str(book["code"])
        book_paragraph = Paragraph(
            f"{escape(str(book['name']))}  |  {escape(str(book['english']))}", book_style
        )
        book_paragraph._bookmark_name = f"book-{code}"
        story.append(book_paragraph)
        chapter_numbers = sorted(chapter for book_code, chapter in chapters if book_code == code)
        for chapter in chapter_numbers:
            story.append(Paragraph(f"{escape(str(book['name']))} {chapter}장", chapter_style))
            body = []
            for item in chapters[(code, chapter)]:
                body.append(
                    f'<font name="BibleKorean" size="6.6" color="#9C6D2D">'
                    f'<super>{item["verse"]}</super></font> {escape(str(item["text"]))}'
                )
            story.append(Paragraph("<br/>".join(body), body_style))

    def first_page(canvas, _document) -> None:
        canvas.saveState()
        canvas.setTitle(TITLE)
        canvas.setAuthor("BibleFrame")
        canvas.setSubject(NOTICE)
        canvas.restoreState()

    def later_pages(canvas, _document) -> None:
        canvas.saveState()
        canvas.setFont("BibleKorean", 8)
        canvas.setFillColor(colors.HexColor("#746F65"))
        canvas.drawString(0.72 * inch, 10.52 * inch, "BibleFrame  |  한국어 가톨릭 성경 73권")
        canvas.drawRightString(7.78 * inch, 0.38 * inch, str(_document.page))
        canvas.restoreState()

    document.multiBuild(story, onFirstPage=first_page, onLaterPages=later_pages)


def write_readme() -> None:
    text = f"""# BibleFrame 전체 본문 다운로드

Public Domain `World English Bible (Catholic)` 73권을 Vertex AI
`Qwen3-Next-80B-A3B-Instruct`로 번역하고 같은 모델로 원문 대조 교정한 한국어 초안입니다.

- 원문: World English Bible (Catholic) · Public Domain
- 한국어 본문·데이터: CC0 1.0 Universal
- 코드: MIT
- 상태: 가톨릭 정경 기반 비공인 기계 번역 초안
- 주의: 한국천주교주교회의 공용 번역본이 아니며 전례용·교리 판정용이 아닙니다.

TXT와 Word/PDF는 읽기·편집용, JSON과 JSONL은 앱·분석·RAG용입니다. PDF에는 73권
자동 목차와 책갈피가 있고, 가독성을 위해 각 절이 새 줄에서 시작합니다.
장 단위 RAG 자료는 `bibleframe-rag.zip`에 별도로 들어 있습니다.
"""
    (DOWNLOADS / "README_DOWNLOADS.md").write_text(text, encoding="utf-8")


def main() -> None:
    DOWNLOADS.mkdir(parents=True, exist_ok=True)
    payload = load_payload()
    write_text(payload)
    write_json(payload)
    write_jsonl(payload)
    write_docx(payload)
    write_pdf(payload)
    write_readme()
    sizes = {
        path.suffix.lstrip(".") or path.name: path.stat().st_size
        for path in DOWNLOADS.glob(f"{STEM}.*")
    }
    print(json.dumps({"downloads": sizes}, ensure_ascii=False))


if __name__ == "__main__":
    main()
