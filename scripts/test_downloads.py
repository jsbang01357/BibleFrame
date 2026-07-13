#!/usr/bin/env python3
"""BibleFrame 전체 본문 다운로드 파일의 내용과 구조를 검사한다."""

from __future__ import annotations

import json
import zipfile
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
DOWNLOADS = ROOT / "site" / "downloads"
STEM = "bibleframe-ko-catholic-73"


def main() -> None:
    txt = (DOWNLOADS / f"{STEM}.txt").read_text(encoding="utf-8")
    assert "# 창세기 (Genesis)" in txt
    assert "# 토빗기 (Tobit)" in txt
    assert "## 요한 복음서 3장" in txt
    assert "16 하느님께서는 세상을" in txt
    assert "Qwen3-Next 80B" in txt
    assert "CC0-1.0" in txt

    payload = json.loads((DOWNLOADS / f"{STEM}.json").read_text(encoding="utf-8"))
    assert payload["meta"]["books"] == 73
    assert payload["meta"]["chapters"] == 1_328
    assert len(payload["verses"]) == 35_379

    lines = (DOWNLOADS / f"{STEM}.jsonl").read_text(encoding="utf-8").splitlines()
    assert len(lines) == 35_379
    first = json.loads(lines[0])
    assert first["id"] == "GEN-1-1"
    assert first["data_license"] == "CC0-1.0"
    assert first["translation_model"] == "Qwen3-Next-80B-A3B-Instruct"

    docx_path = DOWNLOADS / f"{STEM}.docx"
    with zipfile.ZipFile(docx_path) as archive:
        assert archive.testzip() is None
        assert "word/document.xml" in archive.namelist()
    document = Document(docx_path)
    headings = [(paragraph.style.name, paragraph.text) for paragraph in document.paragraphs]
    assert sum(style == "Heading 1" for style, _ in headings) == 73
    assert sum(style == "Heading 2" for style, _ in headings) == 1_328
    assert any(style == "Heading 1" and text.startswith("토빗기") for style, text in headings)
    assert any(style == "Heading 2" and text == "요한 복음서 3장" for style, text in headings)
    assert any("하느님께서는 세상을 이처럼 사랑하셔서" in paragraph.text for paragraph in document.paragraphs)

    pdf_path = DOWNLOADS / f"{STEM}.pdf"
    reader = PdfReader(pdf_path)
    assert 800 < len(reader.pages) < 2_000
    assert reader.metadata.title == "BibleFrame 한국어 가톨릭 성경 73권"
    assert "BibleFrame" in (reader.pages[0].extract_text() or "")
    assert "목차" in (reader.pages[1].extract_text() or "")
    assert "창세기 | Genesis" in (reader.pages[1].extract_text() or "")
    assert "요한 묵시록 | Revelation" in (reader.pages[4].extract_text() or "")
    assert "요한 묵시록 22장" in (reader.pages[-1].extract_text() or "")
    assert len(reader.outline) == 73
    genesis_lines = (reader.pages[5].extract_text() or "").splitlines()
    assert any(line.startswith("1 한처음에 하느님께서") for line in genesis_lines)
    assert any(line.startswith("2 땅은 공허하고") for line in genesis_lines)
    assert any(line.startswith("3 하느님께서 이르셨다") for line in genesis_lines)
    assert pdf_path.stat().st_size < 25_000_000

    readme = (DOWNLOADS / "README_DOWNLOADS.md").read_text(encoding="utf-8")
    assert "Public Domain" in readme and "CC0 1.0 Universal" in readme and "MIT" in readme
    print("OK: TXT · JSON · JSONL · DOCX · PDF 전체 본문 다운로드 검증 완료")


if __name__ == "__main__":
    main()
